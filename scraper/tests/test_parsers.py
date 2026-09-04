from __future__ import annotations

import io

from PIL import Image
from docx import Document
from openpyxl import Workbook
from pypdf import PdfReader, PdfWriter

from app.document_intelligence.parsers.document import parse_bytes
from app.document_intelligence.parsers.errors import ParseError
from app.document_intelligence.parsers.models import PageResult, ParsedDocument
from app.document_intelligence.parsers.pdf import parse_pdf
from app.document_intelligence.parsers.structured import parse_job_description, parse_resume


def test_txt_parser_normalizes_and_preserves_page_number() -> None:
    parsed = parse_bytes(b"  Alice  \r\n\r\nPython   SQL  ", "resume.txt")
    assert parsed.text == "Alice\n\nPython SQL"
    assert parsed.pages[0].page_number == 1
    assert parsed.confidence == 1.0


def test_csv_parser_is_deterministic() -> None:
    parsed = parse_bytes(b"Name,Skill\nAlice,Python\n", "skills.csv")
    assert parsed.text == "Name | Skill\nAlice | Python"


def test_xlsx_parser_reads_each_sheet_as_a_page() -> None:
    workbook = Workbook()
    workbook.active["A1"] = "Role"
    workbook.active["B1"] = "Engineer"
    second = workbook.create_sheet("Skills")
    second["A1"] = "Python"
    output = io.BytesIO()
    workbook.save(output)
    parsed = parse_bytes(output.getvalue(), "profile.xlsx")
    assert len(parsed.pages) == 2
    assert "Role | Engineer" in parsed.text
    assert parsed.pages[1].page_number == 2


def test_docx_parser_extracts_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Summary")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    output = io.BytesIO()
    document.save(output)
    parsed = parse_bytes(output.getvalue(), "resume.docx")
    assert "Summary" in parsed.text
    assert "Skill | Python" in parsed.text


def test_html_parser_removes_executable_content() -> None:
    parsed = parse_bytes(b"<h1>Engineer</h1><script>secret()</script><p>Python</p>", "jd.html")
    assert parsed.text == "Engineer\nPython"
    assert "secret" not in parsed.text


def test_image_parser_reports_ocr_failure_without_inventing_text(monkeypatch) -> None:
    image = Image.new("RGB", (20, 20), "white")
    output = io.BytesIO()
    image.save(output, format="PNG")

    def fail(*_args, **_kwargs):
        raise RuntimeError("tesseract unavailable")

    monkeypatch.setattr("pytesseract.image_to_data", fail)
    try:
        parse_bytes(output.getvalue(), "scan.png")
    except ParseError as error:
        assert error.code == "OCR_FAILED"
        assert error.stage == "ocr"
    else:
        raise AssertionError("OCR failure should be surfaced")


def test_pdf_scanned_page_uses_ocr_and_keeps_page_number(monkeypatch) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    output = io.BytesIO()
    writer.write(output)

    monkeypatch.setattr(
        "app.document_intelligence.parsers.pdf._ocr_page",
        lambda *_args: ("Scanned text", 88.0, []),
    )
    pages, warnings = parse_pdf(output.getvalue())
    assert pages[0].extraction_method == "ocr"
    assert pages[0].text == "Scanned text"
    assert warnings == []


def test_resume_and_jd_structured_parsers_leave_missing_values_empty() -> None:
    resume = ParsedDocument(
        parser_version="test",
        filename="resume.pdf",
        media_type="application/pdf",
        pages=[PageResult(
            page_number=1,
            text="Alice Example\nalice@example.com\n\nSkills\nPython\nExperience\nEngineer",
            extraction_method="text",
        )],
        text="Alice Example\nalice@example.com\n\nSkills\nPython\nExperience\nEngineer",
        confidence=1.0,
        review_required=False,
    )
    resume_result = parse_resume(resume)
    assert resume_result is not None
    assert resume_result.name == "Alice Example"
    assert resume_result.contact_details["email"] == "alice@example.com"
    assert resume_result.skills == ["Python"]

    jd = ParsedDocument(
        parser_version="test",
        filename="jd.txt",
        media_type="text/plain",
        pages=[PageResult(
            page_number=1,
            text="Job Title: Backend Engineer\nResponsibilities\nBuild APIs",
            extraction_method="text",
        )],
        text="Job Title: Backend Engineer\nResponsibilities\nBuild APIs",
        confidence=1.0,
        review_required=False,
    )
    jd_result = parse_job_description(jd)
    assert jd_result is not None
    assert jd_result.job_title == "Backend Engineer"
    assert jd_result.company is None
    assert jd_result.responsibilities == ["Build APIs"]


def test_resume_parses_inline_and_key_skills_aliases() -> None:
    text = (
        "Shabeena Sultana Shaik\n"
        "shabeena@example.com\n\n"
        "Technical Skills: Selenium, Java, API Testing\n\n"
        "Key Skills\nPostman\nCypress\n\n"
        "Experience\nSDET Trainee at QSpiders"
    )
    resume = ParsedDocument(
        parser_version="test",
        filename="resume.txt",
        media_type="text/plain",
        pages=[PageResult(page_number=1, text=text, extraction_method="text")],
        text=text,
        confidence=1.0,
        review_required=False,
    )
    result = parse_resume(resume)
    assert result is not None
    joined = " ".join(result.skills).casefold()
    assert "selenium" in joined
    assert "java" in joined
    assert "api testing" in joined
    assert "postman" in joined
    assert "cypress" in joined
