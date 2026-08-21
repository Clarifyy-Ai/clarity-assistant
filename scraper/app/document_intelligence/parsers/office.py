from __future__ import annotations

import csv
import io
from pathlib import PurePosixPath

from app.document_intelligence.parsers.errors import ParseError
from app.document_intelligence.parsers.models import PageResult
from app.document_intelligence.parsers.normalize import normalize_text
from app.document_intelligence.parsers.text import decode_text


def parse_docx(data: bytes) -> list[PageResult]:
    try:
        from docx import Document
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseError("DOCX_CORRUPT", "DOCX could not be opened.", stage="decoding") from exc
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_references = [f"table-{index}" for index, _ in enumerate(document.tables, start=1)]
    for table in document.tables:
        lines.append(" | ".join(cell.text.strip() for row in table.rows for cell in row.cells))
    return [PageResult(
        page_number=1,
        text=normalize_text("\n".join(lines)),
        extraction_method="text",
        table_references=table_references,
    )]


def parse_csv(data: bytes) -> list[PageResult]:
    try:
        text = decode_text(data)
        rows = csv.reader(io.StringIO(text))
        lines = [" | ".join(cell.strip() for cell in row) for row in rows]
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError("CSV_CORRUPT", "CSV could not be parsed.", stage="decoding") from exc
    return [PageResult(page_number=1, text=normalize_text("\n".join(lines)), extraction_method="text")]


def parse_xlsx(data: bytes) -> list[PageResult]:
    try:
        from openpyxl import load_workbook
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:
        raise ParseError("XLSX_CORRUPT", "XLSX could not be opened.", stage="decoding") from exc
    pages: list[PageResult] = []
    for number, sheet in enumerate(workbook.worksheets, start=1):
        lines = [
            " | ".join("" if value is None else str(value).strip() for value in row)
            for row in sheet.iter_rows(values_only=True)
        ]
        pages.append(PageResult(
            page_number=number,
            text=normalize_text("\n".join(lines)),
            extraction_method="text",
            table_references=[f"sheet-{sheet.title}"],
        ))
    return pages


def parse_office(data: bytes, filename: str) -> list[PageResult]:
    suffix = PurePosixPath(filename.lower()).suffix
    if suffix == ".docx":
        return parse_docx(data)
    if suffix == ".csv":
        return parse_csv(data)
    if suffix == ".xlsx":
        return parse_xlsx(data)
    raise ParseError("UNSUPPORTED_FORMAT", "Office format is not supported.", stage="validation")
